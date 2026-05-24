from crewai import Agent, Task, Crew
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
import os
import streamlit as st
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
import base64
from langchain_xai import ChatXAI

# Carrega as variáveis de ambiente
load_dotenv()

# Configurações de API
SERPER_API_KEY = os.getenv('SERPER_API_KEY')

os.environ['XAI_API_KEY'] = os.getenv('XAI_API_KEY')

# Verificação das variáveis de ambiente
if not SERPER_API_KEY or not os.environ['XAI_API_KEY']:
    raise EnvironmentError("Certifique-se de configurar as variáveis SERPER_API_KEY e XAI_API_KEY no arquivo .env.")

os.environ['SERPER_API_KEY'] = SERPER_API_KEY

def generate_docx(result):
    doc = Document()
    doc.add_heading("Recomendações de diagnóstico e tratamento de saúde", 0)
    # Garante que 'result' seja uma string
    doc.add_paragraph(str(result))
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'

st.set_page_config(
    layout='wide'
)

st.title("IA Assistente Médico")

# Input do usuário
gender = st.selectbox('Selecione o Gênero', ("Masculino", "Feminino", "Outro"))
age = st.number_input("Qual é a idade?", min_value=0, max_value=120, value=25)
symptoms = st.text_area("Quais são os sintomas?", 'e.g., Febre, Tosse, Dor de cabeça')
medical_history = st.text_area("Forneça o histórico médico", 'e.g., Diabetes, Hypertensão')

# Ferramentas
search_tool = SerperDevTool(api_key=SERPER_API_KEY)
scrape_tool = ScrapeWebsiteTool()

# Configuração do LLM com provedor explícito
llm = ChatXAI(
    xai_api_key=os.environ['XAI_API_KEY'],
    model="xai/grok-beta",
    provider="openai"  # Substitua por outro provedor, caso necessário
)

# Agentes
diagnostician = Agent(
    role="Medical Diagnostician",
    goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
    backstory="This agent specializes in diagnosing medical conditions based on patient-reported symptoms and medical history. It uses advanced algorithms and medical knowledge to identify potential health issues.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

treatment_advisor = Agent(
    role="Treatment Advisor",
    goal="Recommend appropriate treatment plans based on the diagnosis provided by the medical Diagnostician.",
    backstory="This agent specializes in creating treatment plans tailored to individual patient needs. It considers the diagnosis, patient history, and current best practices in medicine to recommend effective treatments.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

# Tarefas
diagnose_task = Task(
    description=(
        "1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).\n"
        "2. Provide a preliminary diagnosis with possible conditions based on the provided information.\n"
        "3. Limit the diagnosis to the most likely conditions."
    ),
    expected_output="A preliminary diagnosis with a list of possible conditions.",
    agent=diagnostician
)

treatment_task = Task(
    description=(
        "1. Based on the diagnosis, recommend appropriate treatment plans step by step.\n"
        "2. Consider the patient's medical history ({medical_history}) and current symptoms ({symptoms}).\n"
        "3. Provide detailed treatment recommendations, including medications, lifestyle changes, and follow-up care."
    ),
    expected_output="A comprehensive treatment plan write in brazilian portuguese language tailored to the patient's needs.",
    agent=treatment_advisor
)

crew = Crew(
    agents=[diagnostician, treatment_advisor],
    tasks=[diagnose_task, treatment_task],
    verbose=True
)

# Botão para iniciar o processo
if st.button("Obter Diagnóstico e Plano de Tratamento"):
    with st.spinner("Estou raciocinando, por favor, aguarde..."):
        results = crew.kickoff(inputs={"symptoms": symptoms, "medical_history": medical_history})
        
        # Converte 'results' para string para evitar erros
        results_text = "\n".join([str(item) for item in results]) if isinstance(results, list) else str(results)
        
        st.write(results_text)
        
        docx_file = generate_docx(results_text)

        download_link = get_download_link(docx_file, "diagnosis_and_treatment_plan.docx")
        st.markdown(download_link, unsafe_allow_html=True)
